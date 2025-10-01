import streamlit as st
from docx import Document
from datetime import datetime
from openai import OpenAI

# ---- Configure OpenAI ----
client = OpenAI(api_key="YOUR_OPENAI_API_KEY")  # replace with your key

st.title("Daily Productivity Report Generator")

st.write("Paste your raw description here:")
description = st.text_area("Your description", height=200)

if st.button("Generate Polished Report"):
    # ---- AI Rewrite ----
    prompt = f"""
    You are an assistant that formats daily work logs into a professional report.
    Input: {description}

    Output must include:
    - Development Tasks & Contributions (with start/end times if mentioned).
    - Bugs/Blockers if any.
    - Testing/Documentation/Code Quality if mentioned.
    - Collaboration.
    - Objectives for Tomorrow.
    Make the text formal and concise.
    """

    response = client.chat.completions.create(
        model="gpt-4o-mini",  # small & cheap, can be gpt-4.1 for more power
        messages=[{"role": "user", "content": prompt}]
    )

    polished = response.choices[0].message.content
    st.subheader("AI-Formatted Report")
    st.write(polished)

    # ---- Load Template ----
    template_path = "Daily Productivity Report-Embedded-Template.docx"
    doc = Document(template_path)

    # Fill Engineer Details
    for p in doc.paragraphs:
        if "Name:" in p.text:
            p.text = "Name: Asem Alroomy"
        if "Date:" in p.text:
            p.text = f"Date: {datetime.today().strftime('%Y-%m-%d')}"

    # Just put the AI-polished text into the first task table for now
    tables = doc.tables
    if tables:
        task_table = tables[0]
        row = task_table.add_row().cells
        row[0].text = "1"
        row[1].text = "AI Generated"
        row[2].text = polished
        row[3].text = "09:00"
        row[4].text = "17:00"
        row[5].text = "N/A"
        row[6].text = "Done"
        row[7].text = ""
        row[8].text = "Self"

    output_path = f"Daily_Report_{datetime.today().strftime('%Y-%m-%d')}.docx"
    doc.save(output_path)

    st.success(f"Report generated: {output_path}")
    with open(output_path, "rb") as f:
        st.download_button("Download Report", f, file_name=output_path)
